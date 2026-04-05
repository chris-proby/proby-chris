# -*- coding: utf-8 -*-
"""
transcriptor.py — 사후인터뷰 MD → 회사 브랜딩 적용 녹취록 .docx

MD 형식: 줄 단위로 `AI 인터뷰어00:00` / `참여자00:32` 다음에 본문.

사용:
  python3 transcriptor.py 전수현.md --participant 전수현
  python3 transcriptor.py 전수현.md --brand spoonlabs
  python3 transcriptor.py 전수현.md --project "흥미도 트래커 사후인터뷰" --brand skt

브랜딩 자동감지:
  --brand 없으면 MD 파일 경로의 폴더명을 brands.json keywords와 매칭.

로고 디렉토리:
  기본값: 환경변수 PROBY_LOGO_DIR 또는 이 스크립트 옆 brands.json과 동일 폴더에서
          ../../../../../../../Desktop/proby-sync/02. services/website/assets/logo 탐색.
  직접 지정: --logo-dir /path/to/logos
"""
from __future__ import annotations

import argparse
import json
import os
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 경로 설정 ──────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
BRANDS_JSON = SCRIPT_DIR / "brands.json"

_DEFAULT_LOGO_CANDIDATES = [
    Path(os.environ["PROBY_LOGO_DIR"]) if "PROBY_LOGO_DIR" in os.environ else None,
    Path.home() / "Desktop/proby-sync/02. services/website/assets/logo",
    Path.home() / "proby-chris/proby-chris/02. services/website/assets/logo",
]
DEFAULT_LOGO_DIR = next(
    (p for p in _DEFAULT_LOGO_CANDIDATES if p and p.is_dir()), None
)

SPEAKER_LINE = re.compile(r"^(AI 인터뷰어|참여자)(\d{2}:\d{2})\s*$")


# ─── 브랜드 정보 ─────────────────────────────────────────────────────────────
@dataclass
class BrandInfo:
    name: str = "Proby"
    logo_path: Path | None = None
    accent_hex: str = "F2F2F2"          # AI 행 배경 틴트 기반색
    ai_bg_hex: str = field(init=False)  # accent의 연한 틴트 (자동 계산)

    def __post_init__(self):
        self.ai_bg_hex = _tint(self.accent_hex, 0.88)  # 88% white mix


def _tint(hex_color: str, factor: float) -> str:
    """hex 색상을 흰색 방향으로 factor 비율만큼 밝게 (0=원색, 1=흰색)."""
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    r2 = int(r + (255 - r) * factor)
    g2 = int(g + (255 - g) * factor)
    b2 = int(b + (255 - b) * factor)
    return f"{r2:02X}{g2:02X}{b2:02X}"


def _accent_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def load_brands() -> dict:
    if BRANDS_JSON.exists():
        return json.loads(BRANDS_JSON.read_text(encoding="utf-8"))
    return {}


def detect_brand(md_path: Path, slug: str | None, logo_dir: Path | None) -> BrandInfo:
    """
    slug 지정 시 brands.json에서 조회.
    미지정 시 md_path 경로의 각 폴더명을 keywords와 대조해 자동 감지.
    로고 파일은 logo_dir에서 탐색.
    """
    brands = load_brands()

    # 1) slug 직접 지정
    if slug:
        key = slug.lower()
        if key in brands:
            return _make_brand_info(brands[key], logo_dir)
        # slug가 로고 파일명과 직접 매칭될 수도 있음
        if logo_dir:
            for ext in ("png", "jpg", "jpeg", "svg", "webp"):
                candidate = logo_dir / f"{key}.{ext}"
                if candidate.exists():
                    return BrandInfo(name=slug, logo_path=candidate)
        print(f"⚠  브랜드 슬러그 '{slug}'를 brands.json에서 찾지 못했습니다. 기본값으로 진행합니다.")
        return BrandInfo()

    # 2) 경로 자동 감지
    path_parts = [p.lower() for p in md_path.parts]
    for brand_key, brand_data in brands.items():
        for kw in brand_data.get("keywords", []):
            if any(kw.lower() in part for part in path_parts):
                print(f"✓  브랜드 자동 감지: {brand_data['name']} (키워드 '{kw}')")
                return _make_brand_info(brand_data, logo_dir)

    # 3) 감지 실패 → fallback
    print("ℹ  브랜드를 자동 감지하지 못했습니다. 기본 스타일로 생성합니다.")
    return BrandInfo()


def _make_brand_info(brand_data: dict, logo_dir: Path | None) -> BrandInfo:
    name = brand_data.get("name", "")
    accent = brand_data.get("accent", "F2F2F2")
    logo_path: Path | None = None
    if logo_dir and brand_data.get("logo"):
        candidate = logo_dir / brand_data["logo"]
        if candidate.exists():
            logo_path = candidate
        else:
            print(f"  ⚠ 로고 파일 없음: {candidate}")
    return BrandInfo(name=name, logo_path=logo_path, accent_hex=accent)


# ─── MD 파싱 ─────────────────────────────────────────────────────────────────
def parse_transcript_md(path: Path) -> list[tuple[str, str, str]]:
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    segments: list[tuple[str, str, str]] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        m = SPEAKER_LINE.match(line)
        if m:
            role = m.group(1)
            ts = m.group(2)
            i += 1
            while i < len(lines) and lines[i].strip() == "":
                i += 1
            body: list[str] = []
            while i < len(lines):
                if SPEAKER_LINE.match(lines[i]):
                    break
                body.append(lines[i])
                i += 1
            segments.append((role, ts, "\n".join(body).strip()))
        else:
            i += 1
    return segments


# ─── Word 헬퍼 ───────────────────────────────────────────────────────────────
def add_paragraph_border_bottom(para, color="000000", size=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_tcW(cell, width_cm: float):
    CM_TW = 567
    tcPr = cell._tc.get_or_add_tcPr()
    for ex in list(tcPr):
        if ex.tag == qn("w:tcW"):
            tcPr.remove(ex)
    el = OxmlElement("w:tcW")
    el.set(qn("w:w"), str(int(width_cm * CM_TW)))
    el.set(qn("w:type"), "dxa")
    tcPr.insert(0, el)


# ─── DOCX 빌드 ───────────────────────────────────────────────────────────────
def build_docx(
    segments: list[tuple[str, str, str]],
    participant: str,
    project: str,
    out_path: Path,
    brand: BrandInfo,
) -> None:
    CM_TW = 567

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 헤더: 로고 or 텍스트 ─────────────────────────────────────────────────
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(header_table)
    header_table.columns[0].width = Cm(6)
    header_table.columns[1].width = Cm(13)

    left_cell = header_table.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    logo_para = left_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if brand.logo_path and brand.logo_path.suffix.lower() not in (".svg",):
        # SVG는 python-docx 미지원 → 텍스트 폴백
        try:
            run = logo_para.add_run()
            run.add_picture(str(brand.logo_path), height=Cm(1.4))
        except Exception:
            _add_logo_text(logo_para, brand.name)
    else:
        _add_logo_text(logo_para, brand.name if brand.name != "Proby" else "Proby")

    divider = doc.add_paragraph()
    # 구분선 색상 = accent 색
    accent_r, accent_g, accent_b = _accent_rgb(brand.accent_hex)
    div_color = f"{accent_r:02X}{accent_g:02X}{accent_b:02X}"
    add_paragraph_border_bottom(divider, color=div_color, size=12)
    divider.paragraph_format.space_before = Pt(6)
    divider.paragraph_format.space_after = Pt(0)

    # ── 메타 ─────────────────────────────────────────────────────────────────
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(meta_table)
    meta_data = [
        [("프로젝트", project), ("참여자", participant)],
        [("인터뷰어", "AI 인터뷰어"), ("날짜 / 방식", "— / 비대면")],
    ]
    col_widths = [Cm(3), Cm(7.5), Cm(3), Cm(5.5)]
    for ci, w in enumerate(col_widths):
        for ri in range(2):
            meta_table.cell(ri, ci).width = w
    for ri, row_data in enumerate(meta_data):
        for ci, (label, value) in enumerate(row_data):
            lp = meta_table.cell(ri, ci * 2).paragraphs[0]
            lr = lp.add_run(label)
            lr.font.name = "Arial"; lr.font.size = Pt(7.5); lr.font.bold = True
            lr.font.color.rgb = RGBColor(130, 130, 130)
            lp.paragraph_format.space_before = Pt(4)
            vp = meta_table.cell(ri, ci * 2 + 1).paragraphs[0]
            vr = vp.add_run(value)
            vr.font.name = "Arial"; vr.font.size = Pt(9); vr.font.bold = False
            vr.font.color.rgb = RGBColor(0, 0, 0)
            vp.paragraph_format.space_before = Pt(4)

    meta_divider = doc.add_paragraph()
    add_paragraph_border_bottom(meta_divider, color="CCCCCC", size=4)
    meta_divider.paragraph_format.space_before = Pt(4)
    meta_divider.paragraph_format.space_after = Pt(12)

    # ── 발화 테이블 ──────────────────────────────────────────────────────────
    AI_BG    = brand.ai_bg_hex
    COLOR_TS = (160, 160, 160)
    COLOR_AI_TXT = (30, 30, 30)
    COLOR_PT_TXT = (10, 10, 10)
    L_W, R_W = 2.8, 13.2
    TOTAL_W  = L_W + R_W

    trans_table = doc.add_table(rows=0, cols=2)
    trans_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(trans_table)
    tbl = trans_table._tbl

    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for tag in ("w:tblW", "w:tblLayout"):
        for el in tblPr.findall(qn(tag)): tblPr.remove(el)
    tblW_el = OxmlElement("w:tblW")
    tblW_el.set(qn("w:w"), str(int(TOTAL_W * CM_TW))); tblW_el.set(qn("w:type"), "dxa")
    tblPr.append(tblW_el)
    tblLayout_el = OxmlElement("w:tblLayout"); tblLayout_el.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout_el)
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid"); tbl.append(tblGrid)
    for gc in list(tblGrid): tblGrid.remove(gc)
    for w in (L_W, R_W):
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(int(w * CM_TW))); tblGrid.append(gc)

    def add_turn(speaker: str, ts: str, text_body: str):
        is_ai = speaker == "AI 인터뷰어"
        row = trans_table.add_row()
        lc, rc = row.cells[0], row.cells[1]
        _set_tcW(lc, L_W); _set_tcW(rc, R_W)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if is_ai:
            set_cell_bg(lc, AI_BG); set_cell_bg(rc, AI_BG)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(10)
        lp.paragraph_format.space_after  = Pt(0)
        lp.paragraph_format.left_indent  = Cm(0.3)
        r = lp.add_run("AI" if is_ai else "참여자")
        r.font.name = "Arial"; r.font.size = Pt(8); r.font.bold = True
        r.font.color.rgb = RGBColor(*(accent_r, accent_g, accent_b) if is_ai else (0, 0, 0))
        tp = lc.add_paragraph()
        tp.paragraph_format.space_before = Pt(2)
        tp.paragraph_format.space_after  = Pt(10)
        tp.paragraph_format.left_indent  = Cm(0.3)
        tr2 = tp.add_run(ts)
        tr2.font.name = "Arial"; tr2.font.size = Pt(7.5)
        tr2.font.color.rgb = RGBColor(*COLOR_TS)
        cp = rc.paragraphs[0]
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.space_after  = Pt(10)
        cp.paragraph_format.left_indent  = Cm(0.3)
        cp.paragraph_format.right_indent = Cm(0.3)
        cr2 = cp.add_run(text_body)
        cr2.font.name = "Malgun Gothic"; cr2.font.size = Pt(9.5)
        cr2.font.color.rgb = RGBColor(*(COLOR_AI_TXT if is_ai else COLOR_PT_TXT))

    for speaker, ts, text in segments:
        add_turn(speaker, ts, text)

    # ── 푸터 ─────────────────────────────────────────────────────────────────
    footer_div = doc.add_paragraph()
    add_paragraph_border_bottom(footer_div, color="000000", size=6)
    footer_div.paragraph_format.space_before = Pt(10)
    footer_div.paragraph_format.space_after  = Pt(4)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_text = f"Confidential — {brand.name} × Proby   |   User Insight Team"
    fr = footer_para.add_run(footer_text)
    fr.font.name = "Arial"; fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor(160, 160, 160)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _add_logo_text(para, name: str):
    """로고 이미지 없을 때 텍스트로 대체."""
    lines = name.split() if " " in name else [name]
    r1 = para.add_run("\n".join(lines))
    r1.font.name = "Arial"; r1.font.size = Pt(22); r1.font.bold = True
    r1.font.color.rgb = RGBColor(0, 0, 0)


# ─── CLI ─────────────────────────────────────────────────────────────────────
def main():
    ap = argparse.ArgumentParser(description="MD 원문 → 브랜딩 적용 녹취록 docx")
    ap.add_argument("md_path", type=Path, help="인터뷰 원문 .md")
    ap.add_argument("--participant", default=None, help="참여자 이름 (기본: md stem)")
    ap.add_argument("--project",     default=None, help="프로젝트명 (기본: md 부모폴더 이름)")
    ap.add_argument("--brand",       default=None, help="브랜드 슬러그 (예: spoonlabs, skt)")
    ap.add_argument("--logo-dir",    default=None, type=Path, help="로고 디렉토리 경로 직접 지정")
    ap.add_argument("-o", "--output",default=None, type=Path, help="출력 경로")
    args = ap.parse_args()

    md_path = args.md_path.resolve()
    if not md_path.is_file():
        raise SystemExit(f"파일 없음: {md_path}")

    participant = args.participant or md_path.stem
    project     = args.project or md_path.parent.name
    logo_dir    = args.logo_dir.resolve() if args.logo_dir else DEFAULT_LOGO_DIR
    out_path    = args.output.resolve() if args.output else md_path.parent / f"{md_path.stem}_녹취록.docx"

    brand = detect_brand(md_path, args.brand, logo_dir)
    print(f"브랜드: {brand.name}  |  accent: #{brand.accent_hex}  |  AI배경: #{brand.ai_bg_hex}")
    print(f"로고: {brand.logo_path or '(텍스트 폴백)'}")

    segments = parse_transcript_md(md_path)
    if not segments:
        raise SystemExit("발화를 파싱하지 못했습니다. 형식: 'AI 인터뷰어00:00' / '참여자00:32' 줄 다음에 본문")

    build_docx(segments, participant, project, out_path, brand)
    print(f"저장 완료: {out_path}  ({len(segments)}턴)")


if __name__ == "__main__":
    main()
