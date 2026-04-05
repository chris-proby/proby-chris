#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
녹취록_정리/*.docx 기반으로 가이드라인_2차 질문별 응답 CSV에
한빛나(N) 열을 추가하고 질문별_응답정리.csv 로 저장합니다.

기존 13명 열은 가이드라인_2차_질문별_응답정리_13명_통합.csv 에서 복사합니다.
이민슷(F) → 이민승(F) 헤더 정정.
"""

from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document

BASE = Path(__file__).resolve().parent
CSV_IN = BASE / "가이드라인_2차_질문별_응답정리_13명_통합.csv"
CSV_OUT = BASE / "질문별_응답정리.csv"
DOCX_DIR = BASE / "녹취록_정리"
HANBITNA_DOCX = DOCX_DIR / "한빛나_녹취록.docx"

NA = "**요약:** — (해당 없음 / 조건 비해당)\n\n**인용:** 「(해당 없음)」"
NQ = "**요약:** **(미질문)**\n\n**인용:** 「(전사상 해당 질문 없음)」"


def fmt(summary: str, quote: str) -> str:
    quote = quote.replace("」", "'").replace("「", "'")
    return f"**요약:** {summary}\n\n**인용:** 「{quote}」"


def extract_blocks(docx_path: Path) -> list[tuple[str, str]]:
    d = Document(str(docx_path))
    turns: list[tuple[str, str]] = []
    for ti, t in enumerate(d.tables):
        if ti < 2:
            continue
        for row in t.rows:
            if len(row.cells) < 2:
                continue
            left = "\n".join(p.text.strip() for p in row.cells[0].paragraphs if p.text.strip())
            right = "\n".join(p.text.strip() for p in row.cells[1].paragraphs if p.text.strip())
            if not right:
                continue
            lines = [x.strip() for x in left.split("\n") if x.strip()]
            if not lines:
                continue
            lab = lines[0]
            if lab in ("AI", "인터뷰어"):
                turns.append(("iv", right))
            elif lab == "참여자":
                turns.append(("pt", right))
    blocks: list[tuple[str, str]] = []
    for i in range(len(turns) - 1):
        if turns[i][0] == "iv" and turns[i + 1][0] == "pt":
            blocks.append((turns[i][1], turns[i + 1][1]))
    return blocks


def hanbitna_cells() -> dict[str, str]:
    """한빛나 전사 Q–A 블록 인덱스 기준 수동 매핑."""
    B = extract_blocks(HANBITNA_DOCX)
    # B[0] 인트로 스킵, 이후 B[1].. 가 본문
    def a(*idx: int) -> str:
        parts = [B[i][1].strip() for i in idx if i < len(B)]
        return " ".join(parts)

    def q(*idx: int) -> str:
        parts = [B[i][0].strip() for i in idx if i < len(B)]
        return " ".join(parts)

    out: dict[str, str] = {}

    out["Q1"] = fmt(
        "공식 홈페이지 가입. 빠른 일처리·타 채널 대비 신뢰성을 이유로 선택.",
        f'{a(1)} / {a(2)}',
    )

    out["Q2"] = fmt(
        "무제한 데이터 요금제(명칭 기억 불명). 데이터 무제한·속도 유지 이유로 선택. 네이버 카페 등에서 직접 발품. 의무 유지 조건은 기억 불명. OTT 연동은 홈페이지 글 안내로 이해(영상·이미지 안내 제안). 인터넷·TV·가족 결합 미신청(가격 메리트 체감 낮음). 유료 부가—스팸방지(사후 유료 가입), 해지·이용법 안내는 명확히 수신.",
        f'{a(3)} … {a(6)} / {a(7)}–{a(9)} / {a(10)}–{a(11)} / {a(12)}–{a(14)}',
    )

    out["Q3"] = fmt("타인 추천 의향 7/10.", a(24))
    out["Q4-1"] = fmt("본인에게는 좋아도 타인에게는 다를 수 있어 만점 미부여.", a(25))
    out["Q4-2"] = fmt(
        "이전 통신사 추천 6/10. 구독 할인 부족, 공공·대중교통 무료 와이파이 포인트가 적어 데이터 부담.",
        f'{a(27)} / {a(28)}',
    )
    out["Q5"] = fmt(
        "현 통신사가 공공·대중교통 무료 와이파이 제공으로 데이터 절감 체감이 이전 대비 낫다고 느낌.",
        a(29),
    )
    out["Q6"] = fmt(
        "(Q4-1 불만 연장) 이전 통신사에서 와이파이 포인트 부족·데이터 부담 경험을 구체적으로 언급.",
        a(28),
    )

    out["Q7"] = fmt(
        "개통 당일—홈페이지 가입·개통 신청·개통 전화 후 개통 완료 위주로 기억.",
        a(30),
    )
    out["Q8"] = fmt("2주 내 특별히 불안·짜증 순간은 없었다고 응답.", a(31))
    out["Q9"] = NQ
    out["Q10"] = fmt("가입 후 문자·알림톡은 적당한 수준으로 느꼈다고 응답.", a(32))
    out["Q11"] = fmt("내용 과다·빈도 부담은 없었다고 응답.", a(33))
    out["Q12"] = fmt(
        "초기 필수 알림으로 혜택·데이터 한도·가격 고지를 제시.",
        a(34),
    )
    out["Q13"] = NQ
    out["Q14"] = NQ
    out["Q15"] = fmt("가입 직후 요금·혜택 관련 별도 문의 필요 상황은 없었다고 응답.", a(35))
    out["Q16"] = NQ
    out["Q17"] = NQ

    out["Q18"] = fmt("데이터 다소 사용·무제한·속도 유지 등으로 무제한 요금제 선택.", a(4))
    out["Q19"] = fmt("네이버 카페 등에서 후기·정보 확인 후 본인이 직접 선택.", a(5))
    out["Q20"] = fmt("의무 유지 조건 포함 여부는 기억나지 않는다고 응답.", a(6))
    out["Q21"] = NA
    out["Q22"] = NA
    out["Q23"] = fmt(
        "OTT 연동 절차는 홈페이지 글 안내로 이해 가능. 글 외 영상·이미지 가이드가 있으면 더 쉬웠을 것.",
        f'{a(8)} / {a(9)}',
    )
    out["Q24"] = NA
    out["Q25"] = fmt(
        "스팸방지는 초기 포함 아님, 이후 스팸 피해 우려로 본인 판단 하 유료 가입.",
        a(13),
    )
    out["Q26"] = fmt("가입 시 해지 방법·사용법은 명확히 안내받았다고 응답.", a(14))
    out["Q27"] = NA  # 미사용 독려·해지 유도 — 본 전사 구간에서 미언급
    out["Q28"] = NA
    out["Q29"] = fmt(
        "결합 미신청. 가격 메리트가 본인에게 크게 와닿지 않아 선택하지 않음.",
        f'{a(10)} / {a(11)}',
    )
    out["Q30"] = NA
    out["Q31"] = fmt("결합 미신청 이유—가격적 강점 체감이 크지 않음.", a(11))
    out["Q32"] = NA
    out["Q33"] = NA
    out["Q34"] = NA
    out["Q35"] = NA

    out["Q36"] = fmt(
        "멤버십 전반 관심은 크지 않음(혜택 축소·사 간 차이 체감). OTT·음악 구독은 관심.",
        a(15),
    )
    out["Q37"] = fmt(
        "일반 등급으로 앎. 등급은 사용(지출)에 연동되는 명확한 give-and-take로 공정한 편이라고 봄.",
        f'{a(16)} / {a(17)}',
    )
    out["Q38"] = fmt(
        "승급 ‘예정일’ 안내는 도움 안 될 것 같고, 남은 조건 안내는 도움될 것.",
        a(18),
    )
    out["Q39"] = fmt("장기 고객 기준은 약 5년으로 생각.", a(19))

    out["Q41"] = fmt(
        "가입 후 첫 달 멤버십 혜택 이용. 넷플릭스 할인이 가장 마음에 듦.",
        f'{a(20)} / {a(21)}',
    )
    out["Q42"] = NA
    out["Q43"] = fmt(
        "선착순 혜택은 정보 탐색·노력이 있으므로 불공정하다고 보지 않음.",
        a(22),
    )
    out["Q44"] = fmt(
        "적극 이용을 위해 관심 분야 파악 후 맞춤형 혜택·안내가 필요하다고 봄.",
        a(23),
    )
    out["Q45"] = fmt("(Q44와 인접) 맞춤형 서비스 제안을 재강조.", a(23))

    out["Q47"] = fmt("속도 저하·끊김 이슈는 겪지 않았다고 응답.", a(36))
    out["Q48"] = NQ
    out["Q49"] = NQ
    out["Q50"] = NQ

    out["Q51"] = fmt(
        "첫 달 청구서 관련 특별히 어렵거나 헷갈린 점·요금 구조 불신·안내와 상이 경험은 없다고 응답.",
        f'{a(37)} / {a(38)}',
    )
    out["Q52"] = fmt(
        "청구 전 예상 금액은 가입 전에 이미 알고 만족하에 신청했다고 응답.",
        a(39),
    )
    out["Q53"] = fmt(
        "요금제 비교를 위해 가격·장단점이 한눈에 보이는 비교표·정리 필요. 현재는 글 위주로 가독성이 낮다고 느낌.",
        f'{a(40)} / {a(41)}',
    )
    out["Q54"] = fmt("해지를 고민한 순간은 없었다고 응답.", a(42))

    # 검증: B 인덱스 범위 (스크립트 유지보수용)
    _ = q(1)
    assert len(B) >= 43, f"expected >=43 blocks, got {len(B)}"

    return out


def main() -> None:
    han = hanbitna_cells()

    with CSV_IN.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        fieldnames = list(reader.fieldnames or [])
        # 이민슷 → 이민승
        fieldnames = [re.sub(r"이민슷\(F\)", "이민승(F)", h) for h in fieldnames]
        if "한빛나(N)" not in fieldnames:
            fieldnames.append("한빛나(N)")

        rows_out: list[dict[str, str]] = []
        for row in reader:
            qid = row["ID"]
            row["한빛나(N)"] = han.get(qid, NQ)
            # 헤더 키 정합: 입력 파일에 이민슷이 있으면 키 치환
            if "이민슷(F)" in row:
                row["이민승(F)"] = row.pop("이민슷(F)")
            rows_out.append(row)

    with CSV_OUT.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames, quoting=csv.QUOTE_MINIMAL)
        w.writeheader()
        for r in rows_out:
            w.writerow({k: r.get(k, "") for k in fieldnames})

    print(f"✅ 작성: {CSV_OUT}")
    print(f"   열 수: {len(fieldnames)} (한빛나(N) 추가, 이민승(F) 정정)")
    print(f"   행 수: {len(rows_out)}")


if __name__ == "__main__":
    main()
